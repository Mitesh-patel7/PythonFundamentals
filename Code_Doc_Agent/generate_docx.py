from docx import Document

def generate_docx(results, output_file):

    doc = Document()

    doc.add_heading(
        "Project Documentation Report",
        level=1
    )

    for item in results:

        doc.add_page_break()

        doc.add_heading(
            item["file"],
            level=2
        )

        doc.add_paragraph(
            item["analysis"]
        )

    doc.save(output_file)